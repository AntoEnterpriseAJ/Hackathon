"""Integration test for POST /api/documents/shift-template."""
from __future__ import annotations

import base64
import io
import json

from docx import Document
from fastapi.testclient import TestClient

from main import app


def _docx_bytes(headings_with_body: list[tuple[str, str]]) -> bytes:
    doc = Document()
    for heading, body in headings_with_body:
        h = doc.add_paragraph(heading)
        h.style = doc.styles["Heading 1"]
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_shift_template_round_trips_docx_and_report():
    client = TestClient(app)

    old = _docx_bytes([
        ("1. Date despre program", "Universitatea Transilvania"),
        ("8. Tematica", "Curs 1: Limite"),
    ])
    template = _docx_bytes([
        ("1. Date despre program", ""),
        ("8. Tematica", ""),
    ])

    response = client.post(
        "/api/documents/shift-template",
        files={
            "old_fd": ("old.docx", old, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "new_template": ("tpl.docx", template, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        },
    )

    assert response.status_code == 200, response.text
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    raw_report = response.headers["x-shift-report"]
    report = json.loads(base64.b64decode(raw_report).decode("utf-8"))
    assert any(m["confidence"] == "exact" for m in report["matches"])

    out = Document(io.BytesIO(response.content))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "Universitatea Transilvania" in text
    assert "Curs 1: Limite" in text


def test_shift_template_rejects_invalid_docx():
    client = TestClient(app)
    response = client.post(
        "/api/documents/shift-template",
        files={
            "old_fd": ("old.docx", b"not a docx", "application/octet-stream"),
            "new_template": ("tpl.docx", b"also not a docx", "application/octet-stream"),
        },
    )
    assert response.status_code == 422
