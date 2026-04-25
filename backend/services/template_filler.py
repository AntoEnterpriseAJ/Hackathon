"""Render a new-template .docx populated with content from the old FD."""
from __future__ import annotations

import copy
import io

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

from services.docx_section_extractor import (
    Section,
    TableBlock,
    TextBlock,
    _heading_level,
)
from services.fd_docx_renderer import apply_admin_fields
from services.template_section_mapper import SectionMatch

PLACEHOLDER_TEXT = "Vă rugăm completați aici conform noului format."


def fill_template(
    template_bytes: bytes,
    old_sections: list[Section],
    new_sections: list[Section],
    matches: list[SectionMatch],
    plan_meta: dict,
) -> bytes:
    """Return the bytes of the filled new-template docx."""
    doc = Document(io.BytesIO(template_bytes))

    old_by_id = {s.id: s for s in old_sections}
    matches_by_new_id = {m.new_section_id: m for m in matches}

    body = doc.element.body
    children = list(body.iterchildren())

    # Identify heading paragraph indices in document order.
    heading_indices: list[int] = []
    for idx, child in enumerate(children):
        if child.tag == qn("w:p"):
            if _heading_level(Paragraph(child, doc)) is not None:
                heading_indices.append(idx)

    # For each (heading_index_i, next_heading_index) pair, replace the slot.
    # Iterate from the LAST slot backwards so earlier indices stay valid.
    for slot_idx in range(len(heading_indices) - 1, -1, -1):
        if slot_idx >= len(new_sections):
            continue
        start = heading_indices[slot_idx] + 1
        end = (
            heading_indices[slot_idx + 1]
            if slot_idx + 1 < len(heading_indices)
            else len(children)
        )
        slot_children = children[start:end]

        new_sec = new_sections[slot_idx]
        match = matches_by_new_id.get(new_sec.id)
        old_sec = (
            old_by_id.get(match.old_section_id)
            if (match and match.old_section_id)
            else None
        )

        # Build replacement XML elements.
        replacement_elements = (
            _build_section_elements(doc, old_sec)
            if old_sec is not None
            else _build_placeholder_elements(doc)
        )

        # Skip body-level metadata (e.g. trailing <w:sectPr>) — only paragraphs
        # and tables are real content slots that we may rewrite.
        content_tags = {qn("w:p"), qn("w:tbl")}
        slot_children = [c for c in slot_children if c.tag in content_tags]

        # Anchor on the next non-content sibling if the slot ends at body end,
        # so re-inserted elements land before any trailing <w:sectPr>.
        if end < len(children):
            anchor = children[end]
        else:
            trailing = [c for c in children[start:] if c.tag not in content_tags]
            anchor = trailing[0] if trailing else None

        # Remove existing slot children.
        for stale in slot_children:
            body.remove(stale)
        # Insert replacements before anchor (or append at end).
        for elem in replacement_elements:
            if anchor is not None:
                anchor.addprevious(elem)
            else:
                body.append(elem)

        # Re-snapshot children since the tree mutated.
        children = list(body.iterchildren())

    apply_admin_fields(doc, plan_meta or {})

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _build_section_elements(doc, section: Section) -> list:
    """Return XML elements (paragraphs + tables) for a section's body."""
    elements = []
    for block in section.body:
        if isinstance(block, TextBlock):
            for text in block.paragraphs:
                p = doc.add_paragraph(text)
                elements.append(_detach(doc, p._p))
        elif isinstance(block, TableBlock):
            rows_data: list[list[str]] = []
            if block.headers:
                rows_data.append(block.headers)
            rows_data.extend(block.rows)
            if not rows_data:
                continue
            cols = max(len(r) for r in rows_data)
            t = doc.add_table(rows=len(rows_data), cols=cols)
            for r_idx, row in enumerate(rows_data):
                for c_idx in range(cols):
                    if c_idx < len(row):
                        t.rows[r_idx].cells[c_idx].text = row[c_idx]
            elements.append(_detach(doc, t._tbl))
    return elements


def _build_placeholder_elements(doc) -> list:
    p = doc.add_paragraph()
    run = p.add_run(PLACEHOLDER_TEXT)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return [_detach(doc, p._p)]


def _detach(doc, element):
    """Remove an element from the document body so we can reinsert it elsewhere."""
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    return element
