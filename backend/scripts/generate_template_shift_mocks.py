"""Generate mock old-FD + new-template .docx pair for the Template Shifter.

Run from the repo root:
    .venv\\Scripts\\python.exe backend\\scripts\\generate_template_shift_mocks.py

Outputs (overwrites):
    backend/mock-data/template-shift/old_fd.docx
    backend/mock-data/template-shift/new_template.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT_DIR = Path(__file__).resolve().parents[1] / "mock-data" / "template-shift"


def _add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles[f"Heading {level}"]


def _add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""


# ---------------------------------------------------------------------------
# OLD FD — 2024 layout, plain Romanian section names, populated content.
# ---------------------------------------------------------------------------
def build_old_fd() -> Document:
    doc = Document()

    _add_heading(doc, "1. Date despre program")
    _add_table(doc, [
        ["Instituția de învățământ superior", "Universitatea Transilvania din Brașov"],
        ["Facultatea", "Matematică și Informatică"],
        ["Departamentul", "Matematică și Informatică"],
        ["Domeniul de studii", "Matematică"],
        ["Ciclul de studii", "Licență"],
        ["Programul de studii / Calificarea", "Matematică informatică"],
    ])

    _add_heading(doc, "2. Date despre disciplină")
    _add_table(doc, [
        ["Denumirea disciplinei", "Analiza matematică I"],
        ["Titular activități curs", "Conf. dr. Ion Popescu"],
        ["Titular activități seminar", "Asist. dr. Maria Ionescu"],
        ["Anul de studiu", "I"],
        ["Semestrul", "1"],
        ["Tipul de evaluare", "Examen"],
        ["Regimul disciplinei", "Obligatorie"],
    ])

    _add_heading(doc, "3. Timpul total estimat (ore pe semestru al activităților didactice)")
    _add_table(doc, [
        ["Nr. ore pe săptămână", "4", "din care: curs", "2", "seminar/laborator", "2"],
        ["Nr. ore pe semestru", "56", "din care: curs", "28", "seminar/laborator", "28"],
    ])
    _add_table(doc, [
        ["Distribuția fondului de timp", "", "", "", "", "ore"],
        ["Studiul după manual, suport de curs, bibliografie și notițe", "", "", "", "", "30"],
        ["Documentare suplimentară în bibliotecă, pe platformele electronice de specialitate și pe teren", "", "", "", "", "20"],
        ["Pregătire seminarii/laboratoare, teme, referate, portofolii și eseuri", "", "", "", "", "25"],
        ["Tutoriat", "", "", "", "", "4"],
        ["Examinări", "", "", "", "", "4"],
        ["Alte activități", "", "", "", "", "2"],
        ["Total ore studiu individual", "", "", "", "", "85"],
        ["Total ore pe semestru", "", "", "", "", "141"],
        ["Numărul de credite", "", "", "", "", "5"],
    ])

    _add_heading(doc, "4. Precondiții (acolo unde este cazul)")
    doc.add_paragraph("De curriculum: cunoștințe de matematică din liceu (analiză, algebră).")
    doc.add_paragraph("De competențe: capacitatea de a opera cu mulțimi numerice și funcții elementare.")

    _add_heading(doc, "5. Condiții (acolo unde este cazul)")
    doc.add_paragraph("De desfășurare a cursului: sală cu videoproiector.")
    doc.add_paragraph("De desfășurare a seminarului: sală standard, tablă.")

    _add_heading(doc, "6. Competențele specifice acumulate")
    _add_table(doc, [
        ["Tip competență", "Cod", "Descriere"],
        ["Competențe profesionale", "C1", "Utilizarea riguroasă a conceptelor de limită, continuitate, derivabilitate și integrabilitate."],
        ["Competențe profesionale", "C2", "Aplicarea metodelor analizei reale în rezolvarea de probleme din matematică și informatică."],
        ["Competențe profesionale", "C3", "Modelarea matematică a unor fenomene din științele aplicate folosind calcul diferențial și integral."],
        ["Competențe transversale", "CT1", "Aplicarea normelor deontologice de cercetare și de comunicare academică."],
        ["Competențe transversale", "CT2", "Lucrul autonom și în echipă pentru rezolvarea de probleme deschise."],
        ["Competențe transversale", "CT3", "Identificarea oportunităților de formare continuă și de dezvoltare profesională."],
    ])

    _add_heading(doc, "7. Obiectivele disciplinei")
    doc.add_paragraph("Obiectivul general: însușirea aparatului fundamental al analizei reale.")
    doc.add_paragraph("Obiective specifice: limite și continuitate; derivabilitate; integrabilitate Riemann.")

    _add_heading(doc, "8. Conținuturi")
    _add_heading(doc, "8.1 Tematica activităților de curs", level=2)
    _add_table(doc, [
        ["Săptămâna", "Temă curs", "Subteme acoperite", "Metode didactice", "Bibliografie", "Ore"],
        ["1", "Mulțimi numerice. Funcții elementare.", "R, Q, structuri de ordine; funcții elementare și inverse.", "Prelegere + demonstrații pe tablă", "[1] cap. 1", "2"],
        ["2", "Șiruri reale. Limite.", "Convergență, criteriul Cauchy, limite remarcabile.", "Prelegere interactivă", "[1] cap. 2, [3] cap. 3", "2"],
        ["3", "Limite de funcții.", "Limite laterale, limite la infinit, asimptote.", "Prelegere + studiu de caz", "[1] cap. 3", "2"],
        ["4", "Continuitatea funcțiilor reale.", "Funcții continue pe intervale, teorema lui Bolzano.", "Prelegere + brainstorming", "[1] cap. 3", "2"],
        ["5", "Derivabilitatea. Reguli de derivare.", "Derivate uzuale, derivata funcției compuse, inverse.", "Prelegere + exerciții ghidate", "[1] cap. 4", "2"],
        ["6", "Aplicații ale derivatei.", "Monotonie, extreme, convexitate, studiul funcției.", "Prelegere + lucru pe grup", "[1] cap. 4, [2] cap. 5", "2"],
        ["7", "Diferențiala. Formula lui Taylor.", "Polinom Taylor, restul Lagrange, aproximări.", "Prelegere", "[1] cap. 5", "2"],
        ["8", "Primitive. Integrala nedefinită.", "Metode de integrare, substituție, părți.", "Prelegere + exerciții", "[1] cap. 6", "2"],
        ["9", "Integrala Riemann.", "Sume Riemann, criterii de integrabilitate.", "Prelegere", "[1] cap. 6, [3] cap. 6", "2"],
        ["10", "Aplicații ale integralei Riemann.", "Lungimi, arii, volume; aplicații în fizică.", "Prelegere + studii de caz", "[1] cap. 7", "2"],
        ["11", "Integrale improprii.", "Convergență, criterii, exemple clasice.", "Prelegere", "[1] cap. 7", "2"],
        ["12", "Șiruri și serii de funcții.", "Convergență punctuală vs. uniformă; criteriul Weierstrass.", "Prelegere", "[3] cap. 7", "2"],
        ["13", "Serii de puteri.", "Raza de convergență; serii Taylor și Maclaurin.", "Prelegere + demonstrații", "[3] cap. 8", "2"],
        ["14", "Recapitulare și sinteză.", "Sinteza temelor; pregătire pentru examen.", "Discuție dirijată", "toată bibliografia", "2"],
    ])

    _add_heading(doc, "8.2 Tematica activităților de seminar", level=2)
    _add_table(doc, [
        ["Săptămâna", "Temă seminar", "Tipuri de probleme", "Mod de evaluare", "Ore"],
        ["1", "Mulțimi și funcții elementare.", "Probleme deschise + algoritmice.", "Observație directă", "2"],
        ["2", "Calculul limitelor de șiruri.", "Aplicarea criteriilor; șiruri recurente.", "Test scurt", "2"],
        ["3", "Limite și continuitate.", "Studiu de caz: discontinuități.", "Verificare temă", "2"],
        ["4", "Derivate — calcule.", "Aplicații în optimizare.", "Lucrare scurtă", "2"],
        ["5", "Studiul funcțiilor.", "Construire grafice, identificare extreme.", "Discuții dirijate", "2"],
        ["6", "Polinom Taylor — aplicații.", "Aproximări numerice.", "Verificare temă", "2"],
        ["7", "Primitive — exerciții variate.", "Substituție, părți, fracții raționale.", "Test parțial 1", "2"],
        ["8", "Calculul de integrale Riemann.", "Calcul direct + tehnici.", "Observație directă", "2"],
        ["9", "Aplicații geometrice ale integralei.", "Arii, volume de revoluție.", "Prezentare grup", "2"],
        ["10", "Integrale improprii.", "Studiul convergenței.", "Lucrare scurtă", "2"],
        ["11", "Serii numerice.", "Criterii de convergență.", "Verificare temă", "2"],
        ["12", "Serii de puteri.", "Determinarea razei de convergență.", "Test parțial 2", "2"],
        ["13", "Probleme de sinteză.", "Probleme tip examen.", "Discuție în grup", "2"],
        ["14", "Simulare examen.", "Subiecte din anii anteriori.", "Lucrare cu notă", "2"],
    ])

    _add_heading(doc, "9. Coroborarea conținuturilor disciplinei cu așteptările reprezentanților comunității epistemice, asociațiilor profesionale și angajatorilor reprezentativi din domeniul aferent programului")
    doc.add_paragraph("Conținutul disciplinei este aliniat cu programele similare din universitățile europene și răspunde cerințelor angajatorilor din IT și învățământ.")

    _add_heading(doc, "10. Evaluare")
    _add_table(doc, [
        ["Tip activitate", "Criterii de evaluare", "Metode de evaluare", "Pondere din nota finală (%)"],
        ["Curs — examen final", "Cunoașterea și înțelegerea aparatului analizei reale; capacitatea de a demonstra rezultatele fundamentale.", "Examen scris, 3 ore, fără materiale.", "50"],
        ["Seminar — verificare 1", "Calculul limitelor, derivabilitate, studiul funcțiilor.", "Lucrare scrisă, 90 min.", "15"],
        ["Seminar — verificare 2", "Integrare Riemann, aplicații, integrale improprii.", "Lucrare scrisă, 90 min.", "15"],
        ["Temă de casă", "Rezolvarea unui set de probleme cu prezentare orală.", "Portofoliu + susținere de 10 min.", "10"],
        ["Activitate continuă", "Prezență, implicare, rezolvare la tablă, calitatea răspunsurilor.", "Observație sistematică pe parcursul semestrului.", "10"],
    ])
    doc.add_paragraph("Standard minim de performanță: obținerea notei minime 5 la examenul scris ȘI minim 5 la media verificărilor de seminar; prezența obligatorie la minim 70% din activitățile de seminar pentru intrarea în examen.")

    _add_heading(doc, "11. Bibliografie")
    doc.add_paragraph("1. Trif T., Analiză matematică, Editura Universității Transilvania, Brașov, 2021.")
    doc.add_paragraph("2. Nicolescu M., Analiză matematică, Editura Didactică, București, 2019.")
    doc.add_paragraph("3. Rudin W., Principles of Mathematical Analysis, McGraw-Hill, 1976.")

    _add_heading(doc, "12. Aprobări")
    doc.add_paragraph("Data avizării în departament: 15/06/2024")
    _add_table(doc, [
        ["Titular curs", "Titular seminar"],
        ["Conf. dr. Ion Popescu", "Asist. dr. Maria Ionescu"],
    ])
    _add_table(doc, [
        ["Director departament", "Decanul facultății"],
        ["Conf. dr. Vasile Georgescu", "Prof. dr. Andrei Marinescu"],
    ])

    return doc


# ---------------------------------------------------------------------------
# NEW TEMPLATE — 2026 layout: empty bodies, slightly renamed/reordered
# headings to exercise the fuzzy + LLM mapper.
# ---------------------------------------------------------------------------
def build_new_template() -> Document:
    doc = Document()

    # Same as before — should match EXACT.
    _add_heading(doc, "1. Date despre program")
    doc.add_paragraph("")  # empty slot

    # Renamed (was "Date despre disciplină") — should match FUZZY.
    _add_heading(doc, "2. Informații despre disciplină")
    doc.add_paragraph("")

    # Heavily renamed — needs LLM (or stays as placeholder when Claude is off).
    _add_heading(doc, "3. Buget de timp pe semestru")
    doc.add_paragraph("")

    # Same as before — EXACT.
    _add_heading(doc, "4. Precondiții (acolo unde este cazul)")
    doc.add_paragraph("")

    # Same as before — EXACT.
    _add_heading(doc, "5. Condiții (acolo unde este cazul)")
    doc.add_paragraph("")

    # Slight rename — FUZZY.
    _add_heading(doc, "6. Competențe specifice")
    doc.add_paragraph("")

    # Same — EXACT.
    _add_heading(doc, "7. Obiectivele disciplinei")
    doc.add_paragraph("")

    # Same parent — EXACT; subsections renamed.
    _add_heading(doc, "8. Conținuturi")
    _add_heading(doc, "8.1 Plan de curs", level=2)  # renamed → FUZZY/LLM
    doc.add_paragraph("")
    _add_heading(doc, "8.2 Plan de seminar", level=2)  # renamed → FUZZY/LLM
    doc.add_paragraph("")
    _add_heading(doc, "8.3 Plan de laborator", level=2)  # NEW slot — placeholder
    doc.add_paragraph("")

    # Heavily renamed — LLM target.
    _add_heading(doc, "9. Relația cu mediul socio-economic și academic")
    doc.add_paragraph("")

    # Same — EXACT.
    _add_heading(doc, "10. Evaluare")
    doc.add_paragraph("")

    # Same — EXACT.
    _add_heading(doc, "11. Bibliografie")
    doc.add_paragraph("")

    # Same — EXACT.
    _add_heading(doc, "12. Aprobări")
    doc.add_paragraph("")
    _add_table(doc, [
        ["Titular curs", "Titular seminar/laborator"],
        ["", ""],
    ])
    _add_table(doc, [
        ["", ""],
        ["", ""],
    ])

    return doc


# ---------------------------------------------------------------------------
# NEW TEMPLATE V2 — 2027 redesign: reordered sections, aggressive renames,
# bilingual headings, split/merged sections, brand-new slots. Stress-tests
# every code path: EXACT, FUZZY, LLM, and PLACEHOLDER.
# ---------------------------------------------------------------------------
def build_new_template_v2() -> Document:
    doc = Document()

    # 1. Brand new section — should land as PLACEHOLDER (no equivalent in old FD).
    _add_heading(doc, "1. Sumar executiv al disciplinei")
    doc.add_paragraph("")

    # 2. Reordered: was section 7 in old → "Obiectivele disciplinei". EXACT body, renamed → FUZZY/LLM.
    _add_heading(doc, "2. Scopul și obiectivele cursului")
    doc.add_paragraph("")

    # 3. Was old §2 "Date despre disciplină" — bilingual rename → LLM.
    _add_heading(doc, "3. Course identification / Identificarea disciplinei")
    doc.add_paragraph("")

    # 4. Was old §1 "Date despre program" — moderately renamed → FUZZY.
    _add_heading(doc, "4. Date despre programul de studii")
    doc.add_paragraph("")

    # 5. Was old §3 — heavily renamed → LLM.
    _add_heading(doc, "5. Volum de muncă și distribuție orară")
    doc.add_paragraph("")

    # 6+7. Old §4 + §5 merged conceptually, but we keep them split with new wording.
    _add_heading(doc, "6. Cerințe prealabile (cunoștințe și competențe)")  # → old §4 FUZZY
    doc.add_paragraph("")
    _add_heading(doc, "7. Resurse logistice necesare")  # → old §5 LLM
    doc.add_paragraph("")

    # 8. Was old §6 — completely rephrased → LLM.
    _add_heading(doc, "8. Rezultate ale învățării (learning outcomes)")
    doc.add_paragraph("")

    # 9. New section — PLACEHOLDER.
    _add_heading(doc, "9. Strategii didactice și metode de predare")
    doc.add_paragraph("")

    # 10. Old §8 "Conținuturi", but split into separate top-level sections.
    _add_heading(doc, "10. Conținut tematic — curs")  # → old §8.1 FUZZY/LLM
    doc.add_paragraph("")
    _add_heading(doc, "11. Conținut tematic — seminar")  # → old §8.2 FUZZY/LLM
    doc.add_paragraph("")
    _add_heading(doc, "12. Conținut tematic — laborator")  # NEW → PLACEHOLDER
    doc.add_paragraph("")
    _add_heading(doc, "13. Conținut tematic — proiect")  # NEW → PLACEHOLDER
    doc.add_paragraph("")

    # 14. Old §10 — kept identical → EXACT.
    _add_heading(doc, "14. Evaluare")
    doc.add_paragraph("")

    # 15. New section — PLACEHOLDER.
    _add_heading(doc, "15. Politica privind frauda academică")
    doc.add_paragraph("")

    # 16. Old §9 — heavily reworded + reordered → LLM.
    _add_heading(doc, "16. Alinierea cu piața muncii și mediul academic internațional")
    doc.add_paragraph("")

    # 17. Old §11 — bilingual rename → FUZZY.
    _add_heading(doc, "17. Bibliografie / References")
    doc.add_paragraph("")

    # 18. New section — PLACEHOLDER.
    _add_heading(doc, "18. Resurse online și platforme suport")
    doc.add_paragraph("")

    # 19. Old §12 — kept identical → EXACT, with the trailing approval tables.
    _add_heading(doc, "19. Aprobări")
    doc.add_paragraph("")
    _add_table(doc, [
        ["Titular curs", "Titular seminar/laborator"],
        ["", ""],
    ])
    _add_table(doc, [
        ["", ""],
        ["", ""],
    ])

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    old_path = OUT_DIR / "old_fd.docx"
    new_path = OUT_DIR / "new_template.docx"
    new_v2_path = OUT_DIR / "new_template_v2.docx"

    build_old_fd().save(str(old_path))
    build_new_template().save(str(new_path))
    build_new_template_v2().save(str(new_v2_path))

    print(f"wrote {old_path} ({old_path.stat().st_size} bytes)")
    print(f"wrote {new_path} ({new_path.stat().st_size} bytes)")
    print(f"wrote {new_v2_path} ({new_v2_path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
