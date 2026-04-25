"""Tests for services.template_filler."""
from __future__ import annotations

import io

from docx import Document

from services.docx_section_extractor import extract_sections
from services.template_filler import fill_template
from services.template_section_mapper import SectionMatch


def _docx_from(builder) -> bytes:
    doc = Document()
    builder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _two_section_template(label_a: str, label_b: str) -> bytes:
    def build(doc):
        a = doc.add_paragraph(label_a)
        a.style = doc.styles["Heading 1"]
        doc.add_paragraph("")  # body slot
        b = doc.add_paragraph(label_b)
        b.style = doc.styles["Heading 1"]
        doc.add_paragraph("")
    return _docx_from(build)


def _two_section_old(label_a: str, body_a: str, label_b: str, body_b: str) -> bytes:
    def build(doc):
        a = doc.add_paragraph(label_a)
        a.style = doc.styles["Heading 1"]
        doc.add_paragraph(body_a)
        b = doc.add_paragraph(label_b)
        b.style = doc.styles["Heading 1"]
        doc.add_paragraph(body_b)
    return _docx_from(build)


def test_fill_template_carries_matched_body():
    template = _two_section_template("1. Date despre program", "2. Date despre disciplina")
    old = _two_section_old(
        "1. Date despre program", "Universitatea Transilvania",
        "2. Date despre disciplina", "Analiza matematica I",
    )
    old_sections = extract_sections(old)
    new_sections = extract_sections(template)

    matches = [
        SectionMatch(new_section_id=new_sections[0].id, old_section_id=old_sections[0].id, confidence="exact"),
        SectionMatch(new_section_id=new_sections[1].id, old_section_id=old_sections[1].id, confidence="exact"),
    ]

    out_bytes = fill_template(template, old_sections, new_sections, matches, plan_meta={})

    out = Document(io.BytesIO(out_bytes))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "Universitatea Transilvania" in text
    assert "Analiza matematica I" in text


def test_fill_template_inserts_red_placeholder_when_unmatched():
    template = _two_section_template("1. Date despre program", "9. Sectiune noua")
    old = _two_section_old(
        "1. Date despre program", "Universitatea Transilvania",
        "8. Tematica", "Curs 1",
    )
    old_sections = extract_sections(old)
    new_sections = extract_sections(template)

    matches = [
        SectionMatch(new_section_id=new_sections[0].id, old_section_id=old_sections[0].id, confidence="exact"),
        SectionMatch(new_section_id=new_sections[1].id, old_section_id=None, confidence="placeholder"),
    ]

    out_bytes = fill_template(template, old_sections, new_sections, matches, plan_meta={})

    out = Document(io.BytesIO(out_bytes))
    placeholder = next(
        p for p in out.paragraphs if "Vă rugăm completați" in p.text
    )
    red_runs = [r for r in placeholder.runs if r.font.color and r.font.color.rgb is not None]
    assert red_runs, "placeholder paragraph must have a coloured run"
    assert str(red_runs[0].font.color.rgb).upper() == "FF0000"
