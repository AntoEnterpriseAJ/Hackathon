"""UC 1.4 — Render an FD draft into the official DOCX template.

Loads `backend/templates/fd_template.docx` (built from the pdf2docx-converted
UTCN fișă), fills the cells we have data for (admin block, hours, credits,
selected CP/CT competencies + R.Î. bullets), and returns the document bytes.

Narrative sections (preconditions, course content, bibliography, evaluation
criteria, signatures) are intentionally left blank — per UC 1.4 spec the
professor or AI copilot fills those from the old FD.
"""
from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.table import _Cell

from schemas.fd_draft import FdDraft, SelectedCompetency


TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "fd_template.docx"


def render_fd_docx(
    *,
    draft: FdDraft,
    plan_meta: dict | None = None,
) -> bytes:
    """Fill the template with draft data and return the .docx bytes."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"FD template missing at {TEMPLATE_PATH}. "
            "Run backend/scripts/build_fd_template.py first."
        )

    doc = Document(str(TEMPLATE_PATH))
    plan_meta = plan_meta or {}

    _fill_section_1_program(doc, plan_meta)
    _fill_section_2_discipline(doc, draft)
    _fill_section_3_hours(doc, draft)
    _fill_section_6_competencies(doc, draft)
    _fill_section_12_approvals(doc, plan_meta)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- section fillers ----------

def _fill_section_1_program(doc, plan_meta: dict) -> None:
    """Section 1 — Date despre program (T0, 6x2)."""
    table = _find_table_starting_with(doc, "1.1 Instituția")
    if table is None:
        return
    cycle_raw = (
        plan_meta.get("nivel_calificare")
        or plan_meta.get("ciclul_de_studii")
        or "Licență"
    )
    program = (
        plan_meta.get("programul_de_studii_universitare_de_licenta")
        or plan_meta.get("programul_de_studii")
        or plan_meta.get("program_studii")
        or ""
    )
    forma = plan_meta.get("forma_de_invatamant")
    if program and forma:
        program = f"{program} ({forma})"
    mapping = {
        "1.1": plan_meta.get("universitatea") or "Universitatea Transilvania din Brașov",
        "1.2": plan_meta.get("facultatea") or "",
        "1.3": (
            plan_meta.get("departamentul")
            or plan_meta.get("departament")
            or plan_meta.get("departamentul_coordonator")
            or ""
        ),
        "1.4": (
            plan_meta.get("domeniul_de_licenta")
            or plan_meta.get("domeniu_studii")
            or plan_meta.get("domeniul_fundamental")
            or ""
        ),
        "1.5": _title_ro(str(cycle_raw)),
        "1.6": str(program),
    }
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip()
        for prefix, value in mapping.items():
            if label.startswith(prefix):
                _write_cell(row.cells[1], value)
                break


def _fill_section_2_discipline(doc, draft: FdDraft) -> None:
    """Section 2 — Date despre disciplină (T1, 5x11 with merged cells)."""
    table = _find_table_starting_with(doc, "2.1 Denumirea")
    if table is None:
        return

    # R0: discipline name in cells 3..10 (one merged region)
    if len(table.rows) > 0 and len(table.rows[0].cells) > 3:
        _write_cell(table.rows[0].cells[3], draft.course_name or "")

    # R1, R2: titular cells 6..10 — leave blank for professor
    # R3: col 1 = year, col 4 = semester, col 7 = eval form, col 10 = regime/conținut
    # R4: same as R3 for col 1/4/7, col 10 = regime/obligativitate
    if len(table.rows) > 3 and len(table.rows[3].cells) > 10:
        r3 = table.rows[3].cells
        _write_cell(r3[1], _str(draft.year))
        _write_cell(r3[4], _str(draft.semester))
        _write_cell(r3[7], draft.evaluation_form or "")
        # R3 col 10 = "Conținut" value (DC/DD/DS/DF — derived from categoria_formativa)
        _write_cell(r3[10], _categoria_short(draft.categoria_formativa))

    if len(table.rows) > 4 and len(table.rows[4].cells) > 10:
        r4 = table.rows[4].cells
        # R4 col 10 = "Obligativitate" value (DI / DO / DOpt / DF)
        _write_cell(r4[10], "DI")  # default obligatorie — Plan rarely splits


def _fill_section_3_hours(doc, draft: FdDraft) -> None:
    """Section 3 — Timpul total estimat (T2, 12x7)."""
    table = _find_table_starting_with(doc, "3.1 Număr de ore")
    if table is None:
        return

    # Parse weekly hours like "2C+0S+2L" or "2/0/2/0" (C/S/L/P) into parts.
    weekly = _parse_weekly_hours(draft.weekly_hours)
    total_weekly = sum(weekly.values()) if weekly else None
    curs_w = weekly.get("C")
    sem_lab_w = (weekly.get("S") or 0) + (weekly.get("L") or 0) + (weekly.get("P") or 0)
    if not weekly:
        sem_lab_w = None

    # R0: col 2 = total weekly, col 4 = curs, col 6 = sem/lab/proiect
    r0 = table.rows[0].cells
    _write_cell(r0[2], _str(total_weekly))
    _write_cell(r0[4], _str(curs_w))
    _write_cell(r0[6], _str(sem_lab_w))

    # R1: total ore semester (planul de învățământ) — assume 14 weeks default
    weeks = 14
    total_sem = total_weekly * weeks if total_weekly is not None else draft.total_hours
    curs_sem = curs_w * weeks if curs_w is not None else None
    sem_lab_sem = sem_lab_w * weeks if sem_lab_w is not None else None
    r1 = table.rows[1].cells
    _write_cell(r1[2], _str(total_sem))
    _write_cell(r1[4], _str(curs_sem))
    _write_cell(r1[6], _str(sem_lab_sem))

    # R9: 3.7 Total ore activitate student | R10: 3.8 Total ore semestru | R11: 3.9 Credite
    if len(table.rows) > 11:
        # 3.9 Credite — last useful field we know.
        _write_cell(table.rows[11].cells[1], _str(draft.credits))
        # Leave 3.7/3.8 blank — derived values, professor adjusts.


def _fill_section_6_competencies(doc, draft: FdDraft) -> None:
    """Section 6 — Competențe specifice (T7 = CP, T9 = CT)."""
    cp_table = _find_table_starting_with(doc, "Competențe profesionale")
    ct_table = _find_table_starting_with(doc, "Competențe transversale")

    if cp_table is not None and len(cp_table.rows[0].cells) > 1:
        _write_cell_lines(cp_table.rows[0].cells[1], _format_competency_block(draft.selected_cp))

    if ct_table is not None and len(ct_table.rows[0].cells) > 1:
        _write_cell_lines(ct_table.rows[0].cells[1], _format_competency_block(draft.selected_ct))


def _format_competency_block(entries: list[SelectedCompetency]) -> list[str]:
    if not entries:
        return ["(de selectat de către profesor)"]
    lines: list[str] = []
    for e in entries:
        lines.append(f"{e.code} {e.title}".strip())
        lines.extend(e.ri_bullets)
        lines.append("")
    if lines and lines[-1] == "":
        lines.pop()
    return lines


def _fill_section_12_approvals(doc, plan_meta: dict) -> None:
    """Section 12 — Aprobări: replace dates in the avizat paragraph and pre-fill T21."""
    date_str = _format_date(plan_meta.get("data_aprobarii"))
    director = (
        plan_meta.get("directorul_de_departament")
        or plan_meta.get("director_departament")
        or ""
    )
    decan = plan_meta.get("decanul_facultatii") or ""

    # 1) Update the avizat-paragraph dates if a plan date is available.
    if date_str:
        for paragraph in doc.paragraphs:
            if "Consiliu de departament" in paragraph.text and "data de" in paragraph.text:
                _replace_paragraph_dates(paragraph, date_str)
                break

    # 2) T21 — signature block (2x2). Pre-fill Director departament (and Decan if useful).
    if not (director or decan or date_str):
        return
    t21 = _find_signature_table(doc)
    if t21 is None:
        return
    rows = t21.rows
    if not rows:
        return
    # Layout convention used here:
    #   R0 = "Titular curs" / "Titular seminar/laborator" — left for the professor.
    #   R1C0 = Director departament label+name; R1C1 = Decan facultate label+name.
    if len(rows) >= 2 and len(rows[1].cells) >= 1 and director:
        _write_cell_lines(rows[1].cells[0], ["Director departament", director])
    if len(rows) >= 2 and len(rows[1].cells) >= 2 and decan:
        _write_cell_lines(rows[1].cells[1], ["Decanul facultății", decan])


def _find_signature_table(doc):
    """Return the empty 2x2 signature table that follows section 11."""
    for table in doc.tables:
        if len(table.rows) == 2 and len(table.rows[0].cells) == 2:
            joined = "".join(c.text.strip() for r in table.rows for c in r.cells)
            if not joined:
                return table
    return None


def _replace_paragraph_dates(paragraph, new_date: str) -> None:
    """Replace any dd/mm/yyyy or yyyy-mm-dd dates in a paragraph's runs."""
    pattern = re.compile(r"\b(\d{1,2}/\d{1,2}/\d{4}|\d{4}-\d{1,2}-\d{1,2})\b")
    for run in paragraph.runs:
        if run.text and pattern.search(run.text):
            run.text = pattern.sub(new_date, run.text)


def _format_date(value) -> str:
    if not value:
        return ""
    text = str(value).strip()
    # ISO yyyy-mm-dd → dd/mm/yyyy (matches the template's existing format).
    m = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})$", text)
    if m:
        y, mo, d = m.groups()
        return f"{int(d):02d}/{int(mo):02d}/{y}"
    return text


# ---------- low-level helpers ----------

def _find_table_starting_with(doc, prefix: str):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if prefix in cell.text:
                    return table
    return None


def _write_cell(cell: _Cell, value: str) -> None:
    """Replace cell contents with a single line of text, preserving cell properties."""
    _clear_cell(cell)
    if value:
        cell.paragraphs[0].add_run(value)


def _write_cell_lines(cell: _Cell, lines: list[str]) -> None:
    """Replace cell contents with one paragraph per line."""
    _clear_cell(cell)
    if not lines:
        return
    cell.paragraphs[0].add_run(lines[0])
    for line in lines[1:]:
        cell.add_paragraph(line)


def _clear_cell(cell: _Cell) -> None:
    """Remove all paragraphs from a cell, leaving one empty paragraph (Word requires it)."""
    paragraphs = list(cell.paragraphs)
    # Keep first paragraph as the anchor; clear its runs.
    first = paragraphs[0] if paragraphs else None
    if first is not None:
        for run in list(first.runs):
            run._element.getparent().remove(run._element)
    # Remove remaining paragraphs entirely.
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)


def _str(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _title_ro(text: str) -> str:
    """Capitalize first letter only, preserve diacritics: 'licență' → 'Licență'."""
    text = (text or "").strip()
    return text[:1].upper() + text[1:] if text else ""


def _categoria_short(categoria: str | None) -> str:
    """Map 'Disciplină de specialitate' → 'DS', 'fundamentală' → 'DF', etc."""
    if not categoria:
        return ""
    c = categoria.lower()
    if "fundament" in c:
        return "DF"
    if "domeniu" in c:
        return "DD"
    if "specialitate" in c:
        return "DS"
    if "complementar" in c:
        return "DC"
    return categoria[:4].upper()


def _parse_weekly_hours(weekly: str | None) -> dict[str, int]:
    """Parse '2C+0S+2L' or '2/0/2/0' into {C: 2, S: 0, L: 2, P: 0}."""
    if not weekly:
        return {}
    out: dict[str, int] = {}
    s = str(weekly).strip()
    # Format A: "2C+0S+2L+0P"
    if any(ch in s for ch in "CSLP"):
        import re
        for m in re.finditer(r"(\d+)\s*([CSLP])", s, re.IGNORECASE):
            out[m.group(2).upper()] = int(m.group(1))
        return out
    # Format B: "2/0/2/0" → C/S/L/P
    parts = [p.strip() for p in s.replace("\\", "/").split("/")]
    keys = ["C", "S", "L", "P"]
    for k, p in zip(keys, parts):
        try:
            out[k] = int(p)
        except (ValueError, TypeError):
            pass
    return out
