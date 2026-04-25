"""Build the Fișa Disciplinei DOCX template from the pdf2docx-converted source.

One-time script:
1. Open `docx_outputs/pdf2docx.docx` (3 stitched FDs from the IA PDF).
2. Truncate the body so only the first discipline's section remains.
3. Blank every table cell, keeping only cells whose text matches a known
   "label" pattern (X.Y prefix, "Bibliografie", section markers, etc.).
4. Save as `backend/templates/fd_template.docx`.

The renderer (`fd_docx_renderer.py`) opens this template and fills cells
by anchor-label lookup, so labels MUST stay intact while values are empty.
"""

from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docx_outputs" / "pdf2docx.docx"
OUTPUT = ROOT / "backend" / "templates" / "fd_template.docx"

# Regexes that identify "label" cells we must NOT blank.
LABEL_PATTERNS = [
    re.compile(r"^\s*\d+\.\d+(\.\d+)*\s+\S"),    # "1.1 X", "10.6 Y" (must have dot + content)
    re.compile(r"^\s*F03\.1-PS7\.2"),            # page-footer code
    re.compile(r"Bibliografie", re.IGNORECASE),
    re.compile(r"Distribu(t|ț)ia fondului de timp", re.IGNORECASE),
    re.compile(r"Studiul după manual", re.IGNORECASE),
    re.compile(r"Documentare suplimentar", re.IGNORECASE),
    re.compile(r"Pregătire seminare", re.IGNORECASE),
    re.compile(r"^\s*Tutoriat\s*$", re.IGNORECASE),
    re.compile(r"^\s*Examinări\s*$", re.IGNORECASE),
    re.compile(r"Alte activități", re.IGNORECASE),
    re.compile(r"^\s*Competențe profesionale\s*$", re.IGNORECASE),
    re.compile(r"^\s*Competențe transversale\s*$", re.IGNORECASE),
    re.compile(r"^\s*Tip de activitate\s*$", re.IGNORECASE),
    re.compile(r"^\s*Metode de predare", re.IGNORECASE),
    re.compile(r"^\s*Număr de ore\s*$", re.IGNORECASE),
    re.compile(r"^\s*Observații\s*$", re.IGNORECASE),
    re.compile(r"^\s*din care:", re.IGNORECASE),
    re.compile(r"^\s*Conținut\d?\)\s*$", re.IGNORECASE),
    re.compile(r"^\s*Obligativitate\d?\)\s*$", re.IGNORECASE),
    re.compile(r"^\s*ore\s*$", re.IGNORECASE),
]


def is_label(text: str) -> bool:
    if not text or not text.strip():
        return False
    for pat in LABEL_PATTERNS:
        if pat.search(text):
            return True
    return False


def blank_cell(cell) -> None:
    """Clear every <w:t> text node inside the cell (XML-level wipe).

    `cell.paragraphs[*].runs[*].text = ""` misses nested text in textboxes
    and merged-cell residues, so we wipe at the element level.
    """
    ns_t = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    for t_el in cell._tc.iter(ns_t):
        t_el.text = ""


def find_first_discipline_end(body) -> int:
    """Find the body-element index where the SECOND discipline starts.

    The second discipline begins with a fresh "1.1 Instituția..." table.
    Returns the index to truncate at (exclusive).
    """
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    seen_first = False
    for i, el in enumerate(list(body)):
        if not el.tag.endswith("}tbl"):
            continue
        first_cell = el.find(f".//{ns}tc")
        if first_cell is None:
            continue
        text = "".join(t.text or "" for t in first_cell.iter() if t.tag.endswith("}t"))
        if "1.1" in text and "Institu" in text:
            if seen_first:
                return i
            seen_first = True
    return len(list(body))


def main() -> int:
    if not SOURCE.exists():
        print(f"[ERR] source DOCX not found: {SOURCE}")
        return 1

    print(f"Opening {SOURCE}")
    doc = Document(str(SOURCE))
    body = doc.element.body

    cutoff = find_first_discipline_end(body)
    elements = list(body)
    print(f"First discipline ends at body element {cutoff} (of {len(elements)})")

    # Remove everything from cutoff onward (except keep sectPr at the very end).
    # sectPr must remain as the last child of body.
    sect_pr = None
    last = elements[-1]
    if last.tag.endswith("}sectPr"):
        sect_pr = deepcopy(last)

    for el in elements[cutoff:]:
        body.remove(el)

    if sect_pr is not None and not list(body)[-1].tag.endswith("}sectPr"):
        body.append(sect_pr)

    # Now blank value cells across remaining tables.
    blanked = 0
    kept = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if is_label(cell.text):
                    kept += 1
                else:
                    blank_cell(cell)
                    blanked += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Blanked {blanked} cells, kept {kept} label cells")
    print(f"Saved {OUTPUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
